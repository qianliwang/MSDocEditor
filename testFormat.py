'''For Quiz Generation'''
from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
import os,sys,re

def formatLine(line):
    pre1 = "p";
    pre2 = "slide ch13s";
    parts = line.split(',');
    return "("+pre1+parts[0]+","+pre2+parts[1].strip("\n")+")";

print "-----Extract input parameters, the page number, the slide number, correct answer, question color-----------";

inputFile = open("./input.txt",'r');
inputLines = inputFile.readlines();
formattedResult = map(formatLine,inputLines);

answers = [l.split(',')[2] for l in inputLines];
questionMarks = [l.split(',')[3].strip("\n") for l in inputLines];
answers_num = [ord(char)-96 for char in answers];
#print answers_num;
#print formattedResult;

print "--------------------Create the test.docx-------------------------";

f = open("../Quiz/ch13/Chapter13Test.docx",'r');
d = Document(f);

newDoc = Document();

#run = newDoc.add_paragraph().add_run("To test the highlight color!!");
#font = run.font;
#font.color.rgb = RGBColor(0xff,0x00,0x00);
#print font.highlight;

pattern = "^(\d{1,2})\W{1,2}\w{2,}";
prog = re.compile(pattern);
count = 0;
mark = 0;
for p in d.paragraphs:
    tempMatch = prog.match(p.text);
    if tempMatch:
        tempIndex = int(tempMatch.group(1))-1;
        print tempIndex;
        tempP = newDoc.add_paragraph(p.text[:1]);
        tempQ = tempP.add_run(p.text[1:2]);
        if questionMarks[tempIndex] == 'y':
            tempQ.font.color.rgb = RGBColor(0xff,0xff,0x00);
        elif questionMarks[tempIndex] == 'g':
            tempQ.font.color.rgb = RGBColor(0x00,0xff,0x00);
        elif questionMarks[tempIndex] == 'r':
            tempQ.font.color.rgb = RGBColor(0xff,0x00,0x00);
        elif questionMarks[tempIndex] == 'p':
            tempQ.font.color.rgb = RGBColor(0x80,0x00,0x80);
        else:
            pass;
        tempP.add_run(p.text[2:]+formattedResult[tempIndex]+"\n");
        mark = count + answers_num[tempIndex];
    else:
        tempAnswer = newDoc.add_paragraph().add_run("\t"+p.text);
        if count == mark:
            tempAnswer.font.color.rgb = RGBColor(0xff,0x00,0x00);
            tempAnswer.bold = True;
    count = count + 1;

sections = newDoc.sections;
section = sections[0];
section.left_margin = Inches(0.7);
section.right_margin = Inches(0.7);
section.top_margin = Inches(0.7);
section.bottom_margin = Inches(0.7);
newDoc.save("test.docx");
f.close();
