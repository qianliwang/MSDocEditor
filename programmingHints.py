'''For ProgrammingHints'''
import os,sys
from docx import Document
'''
To extract the content of the html and css files and put into a MS doc file. Then to edit the doc file and add programming
hints into it. The programming hint is related to the grading scale, and just added into the doc file.
'''
def getHTMLandCSSFile(folder):
    folderList = os.listdir(folder);
    fileList = [];
    for f in folderList:
        if f.endswith("html"):
            fileList.append(f);
        elif f.endswith("css"):
            fileList.append(f);
        else:
            pass;
    return fileList;

folderPath = "../programmingHints/chapter14jQuery/";
htmlFileList = getHTMLandCSSFile(folderPath);

newDoc = Document();

count = 0;
for f in htmlFileList:
    newDoc.add_heading(f,level=1);
    file = open(folderPath+"/"+f,'r');
    lines = file.readlines();
    for line in lines:
        newDoc.add_paragraph(line.strip("\n\r"));

newDoc.save("test_hints.docx");
