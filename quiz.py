from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
import os,sys,re

'''
Each chapter has a quiz set. The procedure is to review all avaiable problems in each chapter. For each problem, recording 
the related textbook pape number, the number of slide of this chapter, the correct answer(Single selection, A,B,C,D), 
and the level of the problem represented by different colors(hard,medium,easy). Then these parameters are
put into a txt file. Later, based on this txt input file, the related page number and the slide number are added into
the head of each problem, the correct answer would be bolded, and the level of each problem is marked with different color.
Finally, output the modified doc file. The second task is to extract certain number of problems as the condidate quiz
problems and saved as another doc file. 
'''
def formatLine(line):
    pre1 = "p";
    pre2 = "slide ch8s";
    parts = line.split(',');
    return "("+pre1+parts[0]+","+pre2+parts[1].strip("\n")+")";

print "-----Extract input parameters, the page number, the slide number, correct answer, question color-----------";

inputFile = open("./input.txt",'r');
inputLines = inputFile.readlines();
formattedResult = map(formatLine,inputLines);

answers = [l.split(',')[2] for l in inputLines];
questionMarks = [l.split(',')[3].strip("\n") for l in inputLines];
answers_num = [ord(char)-96 for char in answers];
#print answers_num;
#print formattedResult;

print "--------------------Create the test.docx-------------------------";

f = open("../Quiz/ch8/Chapter8Test.docx",'r');
d = Document(f);

newDoc = Document();

#run = newDoc.add_paragraph().add_run("To test the highlight color!!");
#font = run.font;
#font.color.rgb = RGBColor(0xff,0x00,0x00);
#print font.highlight;

pattern = "^(\d{1,2})\W{1,2}\w{2,}";
prog = re.compile(pattern);
count = 0;
mark = 0;
for p in d.paragraphs:
    tempMatch = prog.match(p.text);
    if tempMatch:
        tempIndex = int(tempMatch.group(1))-1;
        print tempIndex;
        tempP = newDoc.add_paragraph(p.text[:1]);
        tempQ = tempP.add_run(p.text[1:2]);
        if questionMarks[tempIndex] == 'y':
            tempQ.font.color.rgb = RGBColor(0xff,0xff,0x00);
        elif questionMarks[tempIndex] == 'g':
            tempQ.font.color.rgb = RGBColor(0x00,0xff,0x00);
        elif questionMarks[tempIndex] == 'r':
            tempQ.font.color.rgb = RGBColor(0xff,0x00,0x00);
        elif questionMarks[tempIndex] == 'p':
            tempQ.font.color.rgb = RGBColor(0x80,0x00,0x80);
        else:
            pass;
        tempP.add_run(p.text[2:]+formattedResult[tempIndex]+"\n");
        mark = count + answers_num[tempIndex];
    else:
        tempAnswer = newDoc.add_paragraph().add_run("\t"+p.text);
        if count == mark:
            tempAnswer.font.color.rgb = RGBColor(0xff,0x00,0x00);
            tempAnswer.bold = True;
    count = count + 1;

sections = newDoc.sections;
section = sections[0];
section.left_margin = Inches(0.7);
section.right_margin = Inches(0.7);
section.top_margin = Inches(0.7);
section.bottom_margin = Inches(0.7);
newDoc.save("test.docx");
f.close();

print "--------------------Create the test_set.docx-------------------------";

f = open("./test.docx",'r');
d = Document(f);
setDoc = Document();
candidateQuestionIndex = [];
count = 1;
for q in questionMarks:
    if q=='g' or q=='y' or q=='r':
        tempList = [];
        tempList.append(count);
        tempList.append(q);
        candidateQuestionIndex.append(tempList);
    count += 1;

#print candidateQuestionIndex;
print "--------------------Save test_set.docx----------------------------";
count = 0;
tempPointer = 0;
for p in d.paragraphs:
    tempMatch = prog.match(p.text);
    if tempMatch:
#        print p.text;
        tempIndex = int(tempMatch.group(1));
        for sp in candidateQuestionIndex:
            if tempIndex==sp[0]:
                tempPointer = count;
                tempP = setDoc.add_paragraph(p.text[:1]);
                tempQ = tempP.add_run(p.text[1:2]);
                if sp[1] == 'y':
                    tempQ.font.color.rgb = RGBColor(0xff,0xff,0x00);
                elif sp[1] == 'g':
                    tempQ.font.color.rgb = RGBColor(0x00,0xff,0x00);
                elif sp[1] == 'r':
                    tempQ.font.color.rgb = RGBColor(0xff,0x00,0x00);
                else:
                    pass;
                tempP.add_run(p.text[2:]);
                mark = count + answers_num[tempIndex-1];
    else:
        if tempPointer <> 0 and count>tempPointer and count <=(tempPointer+4):
            tempAnswer = setDoc.add_paragraph().add_run("\t"+p.text);
            if count == mark:
                tempAnswer.font.color.rgb = RGBColor(0xff,0x00,0x00);
                tempAnswer.bold = True;
    count = count + 1;
f.close();

sections = setDoc.sections;
section = sections[0];
section.left_margin = Inches(0.7);
section.right_margin = Inches(0.7);
section.top_margin = Inches(0.7);
section.bottom_margin = Inches(0.7);

setDoc.save("test_set.docx");
print "--------------------save test_set.docx-----------------------------";
